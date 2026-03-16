import os
import ssl
import uuid
import re
import gc
import math
from pathlib import Path
from groq import Groq
from dotenv import load_dotenv
from moviepy import VideoFileClip
from pydub import AudioSegment
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from flask import Flask, request, jsonify, send_file

ssl._create_default_https_context = ssl._create_unverified_context
load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

VIDEO_EXTENSIONS = ('.mp4', '.mkv', '.mov', '.avi', '.flv', '.wmv')
AUDIO_EXTENSIONS = ('.mp3', '.wav', '.m4a', '.flac')

GROQ_MAX_BYTES = 20 * 1024 * 1024  # 20MB safe limit (Groq hard limit is 25MB)

app = Flask(__name__)


# ─── Core Functions ────────────────────────────────────────────────

def extract_audio(video_path: str) -> str:
    """Extract audio from video file."""
    audio_path = str(UPLOAD_DIR / f"{uuid.uuid4().hex}.wav")
    video = VideoFileClip(video_path)
    try:
        video.audio.write_audiofile(audio_path, codec='pcm_s16le', fps=16000, logger=None)
    finally:
        video.close()
        gc.collect()
    return audio_path


def split_audio(file_path: str) -> list:
    """
    Split audio into chunks under GROQ_MAX_BYTES.
    Returns list of chunk file paths.
    """
    file_size = os.path.getsize(file_path)

    # No splitting needed
    if file_size <= GROQ_MAX_BYTES:
        return [file_path]

    ext = Path(file_path).suffix.lower().strip('.')
    fmt = ext if ext in ('mp3', 'wav', 'flac', 'm4a') else 'wav'

    audio = AudioSegment.from_file(file_path, format=fmt)

    # Calculate how many chunks we need
    num_chunks = math.ceil(file_size / GROQ_MAX_BYTES)
    chunk_duration_ms = math.ceil(len(audio) / num_chunks)

    chunk_paths = []
    for idx in range(num_chunks):
        start = idx * chunk_duration_ms
        end = min((idx + 1) * chunk_duration_ms, len(audio))
        chunk = audio[start:end]

        chunk_path = str(UPLOAD_DIR / f"{uuid.uuid4().hex}_chunk{idx}.wav")
        chunk.export(chunk_path, format="wav")
        chunk_paths.append(chunk_path)

    return chunk_paths


def transcribe_audio(file_path: str) -> str:
    """
    Transcribe audio using Groq Whisper API.
    Automatically splits large files into chunks under 20MB.
    """
    client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    chunk_paths = split_audio(file_path)
    full_transcription = []

    try:
        for chunk_path in chunk_paths:
            with open(chunk_path, "rb") as f:
                result = client.audio.transcriptions.create(
                    file=(os.path.basename(chunk_path), f),
                    model="whisper-large-v3",
                    response_format="text"
                )
            text = result if isinstance(result, str) else getattr(result, "text", "")
            if text:
                full_transcription.append(text.strip())
    finally:
        # Clean up chunk files (but not the original)
        for chunk_path in chunk_paths:
            if chunk_path != file_path and os.path.exists(chunk_path):
                os.remove(chunk_path)

    return " ".join(full_transcription)


def analyze_with_groq(text: str, instructions: str = "") -> str:
    client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    safe_text = text[:15000]

    system_message = """You are a senior GHL (GoHighLevel) Technical Architect with deep expertise in CRM automation, funnel building, workflow design, and AI-powered client engagement systems.

Your role is to produce professional, implementation-ready Technical Approach Documents that a development team can directly follow. Write with precision, use proper markdown formatting, and ensure every recommendation is actionable and specific to the client's needs."""

    if safe_text.strip():
        source_context = "meeting transcription"
        source_rule = "- Base the document STRICTLY on what is discussed or implied in the transcription. Do NOT fabricate details."
    else:
        source_context = "instructions/notes provided"
        source_rule = "- Base the document STRICTLY on the instructions/notes provided. Use them as the project brief and build a complete technical approach around them."

    prompt = f"""Analyze the following {source_context} and generate a comprehensive **TECHNICAL APPROACH DOCUMENT**.

**RULES:**
{source_rule}
- OMIT any section that is NOT relevant to the {source_context} (e.g., skip Chatbots if never mentioned).
- Use professional markdown formatting with headers, tables, and bullet points.
- Be specific with field names, pipeline stage names, workflow triggers, and automation logic.
- Where the transcription is vague or details are missing, DO NOT ask the client for information. Instead, suggest the most appropriate values, configurations, or approaches based on industry best practices and context clues from the transcription. Clearly mark these as "Suggested:" so the team knows they are recommendations rather than confirmed requirements.
- Provide FULL DETAILED step-by-step breakdowns for every workflow, pipeline, funnel, and website section. Do not summarize — list every individual step, trigger, action, condition, and outcome.
- For any website or landing page, specify whether it should be built using GHL native builder or custom HTML, and detail every section/block of the page.
- Each module must include a numbered step-by-step implementation guide.
- Include a table of all third-party tools/integrations required with their purpose, cost (if known), and integration method.
- Include a Prerequisites section listing everything needed before implementation can begin.

---

**DOCUMENT STRUCTURE** (include only relevant sections):

# Technical Approach : [Client Name or Company Name]
(Extract client/company name from the transcription or suggest based on context.)

# 1. PROJECT OBJECTIVES
- Summarize the business problem the client wants to solve.
- List 3-5 specific, measurable objectives discussed in the meeting.
- Identify the target audience or customer segment if mentioned.

# 2. FEATURE LISTING
| # | Feature/Module | Description | Build Method |
|---|---|---|---|
| 1 | (e.g., Lead Capture Funnel) | (brief description) | (GHL Native / Custom HTML / API) |

# 3. FEATURE DETAILING

## 3.1 Lead Qualification & Custom Forms
## 3.2 Funnels & Landing Pages
## 3.3 Website (if applicable)
## 3.4 Calendar & Scheduling
## 3.5 Pipeline & Deal Tracking
## 3.6 Proposals, Contracts & Payments
## 3.7 Automation Workflows
## 3.8 AI & Chatbot Configuration

# 4. THIRD-PARTY INTEGRATIONS & TOOLS
| Tool/Service | Purpose | Integration Method | Required Plan/Cost | Priority |
|---|---|---|---|---|

# 5. DASHBOARD & REPORTING
# 6. AGENT / TEAM MANUAL RESPONSIBILITIES
# 7. SUGGESTIONS & ASSUMPTIONS
# 8. PREREQUISITES
# 9. NEXT STEPS
- "Please provide GHL sub-account access to dev.patel@theonetechnologies.co.in to begin implementation."

---

{f"**ADDITIONAL INSTRUCTIONS FROM THE TEAM:**" + chr(10) + instructions + chr(10) + "---" + chr(10) if instructions.strip() else ""}
{f"**TRANSCRIPTION:**" + chr(10) + safe_text if safe_text.strip() else "**NOTE:** No transcription provided. Generate the document based entirely on the additional instructions above."}
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
        max_tokens=8000,
        top_p=0.9,
    )
    return response.choices[0].message.content


def add_formatted_runs(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif '*' in part:
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sp in sub_parts:
                if not sp:
                    continue
                if sp.startswith('*') and sp.endswith('*') and len(sp) > 2:
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sp)
        else:
            paragraph.add_run(part)

    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def flush_table(doc, table_rows):
    if not table_rows:
        return
    num_cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for ri, row_data in enumerate(table_rows):
        for ci in range(num_cols):
            cell = table.rows[ri].cells[ci]
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_runs(p, cell_text)

            if ri == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'B92328', qn('w:val'): 'clear'
                })
                shading.append(shading_elm)
            else:
                for run in p.runs:
                    run.font.size = Pt(10)
                if ri % 2 == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): 'F5F5F5', qn('w:val'): 'clear'
                    })
                    shading.append(shading_elm)

    doc.add_paragraph()


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
        s = doc.styles[level]
        s.font.color.rgb = RGBColor(0xB9, 0x23, 0x28)
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(18 if level == 'Heading 1' else 12)
        s.paragraph_format.space_after = Pt(6)

    lines = markdown_text.split('\n')
    i = 0
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(set(c) <= set('-: ') for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        if not stripped:
            i += 1
            continue

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip().strip('*#').strip(), level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip().strip('*#').strip(), level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip().strip('*#').strip(), level=1)
        elif stripped.startswith('---'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('_' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
        elif re.match(r'^(\s{2,})([-*])\s', line):
            text = re.sub(r'^\s+[-*]\s', '', line)
            p = doc.add_paragraph(style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(2)
            add_formatted_runs(p, text)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, stripped[2:])
        elif re.match(r'^\d+[\.\)]\s', stripped):
            match = re.match(r'^(\d+)[\.\)]\s*(.*)', stripped)
            if match:
                p = doc.add_paragraph(style='List Number')
                add_formatted_runs(p, match.group(2))
        else:
            p = doc.add_paragraph()
            add_formatted_runs(p, stripped)

        i += 1

    if in_table and table_rows:
        flush_table(doc, table_rows)

    doc.save(output_path)
    return output_path


# ─── Flask API Routes ───────────────────────────────────────────────

@app.route("/api/generate", methods=["POST"])
def generate():
    file = request.files.get("file")
    instructions = request.form.get("instructions", "").strip()

    has_file = file is not None and file.filename
    if not has_file and not instructions:
        return jsonify({"error": "Please provide either a file or instructions."}), 400

    file_id = uuid.uuid4().hex
    upload_path = None
    temp_audio = None

    try:
        transcription = ""

        if has_file:
            file_ext = os.path.splitext(file.filename)[1].lower()

            if file_ext not in VIDEO_EXTENSIONS + AUDIO_EXTENSIONS:
                return jsonify({
                    "error": f"Unsupported file type: {file_ext}",
                    "supported_video": list(VIDEO_EXTENSIONS),
                    "supported_audio": list(AUDIO_EXTENSIONS)
                }), 400

            upload_path = str(UPLOAD_DIR / f"{file_id}{file_ext}")
            file.save(upload_path)

            if file_ext in VIDEO_EXTENSIONS:
                audio_path = extract_audio(upload_path)
                temp_audio = audio_path
            else:
                audio_path = upload_path

            transcription = transcribe_audio(audio_path)

        analysis = analyze_with_groq(transcription, instructions)

        docx_filename = f"GHL_Solution_Architect_{file_id}.docx"
        docx_path = str(UPLOAD_DIR / docx_filename)
        markdown_to_docx(analysis, docx_path)

        return send_file(
            docx_path,
            as_attachment=True,
            download_name="GHL_Solution_Architect.docx",
            mimetype="application/octet-stream"
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        if upload_path and os.path.exists(upload_path):
            os.remove(upload_path)
        if temp_audio and os.path.exists(temp_audio):
            os.remove(temp_audio)


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "service": "GHL Solution Architect API"})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)